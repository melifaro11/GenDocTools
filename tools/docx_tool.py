from io import BytesIO
from json import dumps
from pathlib import Path
from docx import Document
from utils.logger import get_logger

from utils.download_file import download_file
from utils.upload_file import upload_file
from utils.knowledge import create_knowledge
from utils.get_user_id import get_user_id
from utils.authorization import _get_bearer_token

logger = get_logger(__name__)

def full_context_docx(file_id, file_name, request, URL):
    """
    Return the structure of a DOCX document including index, style, and text of each element.

    Returns:
        str: JSON-formatted string with the document structure or an error object.
    """
    bearer_token = _get_bearer_token(request)

    try:
        docx_file = download_file(
            url=URL,
            token=bearer_token,
            file_id=file_id
        )

        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(docx_file, indent=4, ensure_ascii=False)

        doc = Document(docx_file)
        text_body = {
            "file_name": file_name,
            "file_id": file_id,
            "body": []
        }

        for idx, part in enumerate(doc.paragraphs):
            text = part.text.strip()
            if not text:
                continue  # Skip empty paragraphs

            style = part.style.name
            text_body["body"].append({
                "index": idx,
                "style": style,
                "text": text
            })

        return dumps(text_body, indent=4, ensure_ascii=False)

    except Exception as e:
        return dumps({"error": {"message": str(e)}}, indent=4, ensure_ascii=False)

def review_docx(file_id, file_name, review_comments, request, URL, ENABLE_CREATE_KNOWLEDGE, REVIEWER_AI_ASSISTANT_NAME):
    """
    Review an existing DOCX document and add comments to specified elements.

    Returns:
        dict or str: Upload response dict (with 'file_path_download') or JSON-formatted error string.
    """
    bearer_token = _get_bearer_token(request)

    user_id = get_user_id(URL, bearer_token)
    if not user_id:
        return dumps({"error": {"message": "Error obtaining user id from token"}}, indent=4, ensure_ascii=False)

    try:
        docx_file = download_file(URL, bearer_token, file_id)
        if isinstance(docx_file, dict) and "error" in docx_file:
            return dumps(docx_file, indent=4, ensure_ascii=False)

        doc = Document(docx_file)
        paragraphs = list(doc.paragraphs)
        for item in review_comments:
            try:
                index = item.index
                comment_text = item.comment
            except (AttributeError, TypeError):
                continue

            if index is None or comment_text is None:
                continue

            if 0 <= index < len(paragraphs):
                para = paragraphs[index]
                if para.runs:
                    doc.add_comment(
                        runs=[para.runs[0]],
                        text=comment_text,
                        author=REVIEWER_AI_ASSISTANT_NAME
                    )

        buffer = BytesIO()
        buffer.name = f'{Path(file_name).stem}_reviewed.docx'
        doc.save(buffer)
        buffer.seek(0)

        response, request_data = upload_file(
            url=URL,
            token=bearer_token,
            file_data=buffer,
            filename=f"{Path(file_name).stem}_reviewed",
            file_type="docx"
        )

        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            create_knowledge_status = create_knowledge(
                url=URL,
                token=bearer_token,
                file_id=request_data['id'],
                user_id=user_id,
                knowledge_name="Documents Reviewed by AI"
            )
            if create_knowledge_status:
                logger.info("=> Knowledge base updated successfully.")
            else:
                logger.error("=> Error creating or updating knowledge base")
        elif "error" in response:
            logger.error("=> Error uploading the file.")
        else:
            logger.info("=> Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response

    except Exception as e:
        return dumps({"error": {"message": str(e)}}, indent=4, ensure_ascii=False)

def generate_word_from_template(doc_metadata, columns_body, doc_dict, file_name, request, URL, ENABLE_CREATE_KNOWLEDGE):
    """
    Generate a Word document from metadata and a list of sections.

    Returns:
        dict: Contains 'file_path_download' with a markdown hyperlink for downloading the generated Word file.
    """
    try:
        # Convert Pydantic models to dicts
        metadata_data = doc_metadata.model_dump()
        sections_data = [section if isinstance(section, dict) else section.model_dump(exclude_none=True) for section in doc_dict]
        
        # Create full doc dict
        doc_full = {
            "metadata": metadata_data,
            "sections": sections_data,
            "font": "Times New Roman",  # Default font
            "columns_body": columns_body
        }
        
        # Create buffer
        buffer = BytesIO()
        buffer.name = f'{file_name}.docx'
        
        # Build the document
        from utils.document_builder import build_docx_from_dict
        buffer = build_docx_from_dict(doc_full, buffer, request, URL)
        
        # Upload the file
        bearer_token = _get_bearer_token(request) 
        if bearer_token:
            logger.info("=> Received authorization header")
        else:
            logger.debug("=> Authorization header not present")

        user_id = get_user_id(URL, bearer_token)
        if not user_id:
            return dumps({"error": {"message": "Error obtaining user id from token"}}, indent=4, ensure_ascii=False)

        response, request_data = upload_file(
            url=URL, 
            token=bearer_token, 
            file_data=buffer,
            filename=file_name,
            file_type="docx"
        )

        if "file_path_download" in response and ENABLE_CREATE_KNOWLEDGE:
            create_knowledge_status = create_knowledge(
                url=URL,
                token=bearer_token,
                file_id=request_data['id'], 
                user_id=user_id
            )
            if create_knowledge_status:
                logger.info("=> Knowledge base updated successfully.")
            else:
                logger.error("=> Error creating or updating knowledge base")
        elif "error" in response:
            logger.error("=> Error uploading the file.")
        else:
            logger.info("=> Skipping knowledge creation because ENABLE_CREATE_KNOWLEDGE is false")

        return response

    except Exception as e:
        logger.error("=> An unknown error occurred during DOCX document generation.")
        return dumps({"error": {"message": str(e)}}, indent=4, ensure_ascii=False)
