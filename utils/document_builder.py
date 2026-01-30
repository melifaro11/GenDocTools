from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import math2docx
import os
import numpy as np
from utils.download_file import download_file
from utils.authorization import _get_bearer_token
from utils.img_dimensions import img_dimensions
import logging
logging.basicConfig(level=logging.INFO, force=True)
logger = logging.getLogger("GenFilesMCP")

def vertical_center(metadata_dict: dict, doc: Document) -> int:
    # Constantes
    EMU_INCH = 914400
    CHAR_WIDTH_FACTOR = 0.7  # Ancho promedio de cada carácter como fracción del tamaño de la fuente
    DEFAULT_LINE_HEIGHT = 0.1667  # Altura de línea en pulgadas (12 pt)

    # Tamaños de fuente por clave (en pulgadas)
    font_sizes = {
        'title': 0.5,
        'subtitle': DEFAULT_LINE_HEIGHT,
        'description': DEFAULT_LINE_HEIGHT,
        'author': DEFAULT_LINE_HEIGHT,
        'month': DEFAULT_LINE_HEIGHT,
        'year': DEFAULT_LINE_HEIGHT,
    }

    # Dimensiones de página y márgenes (en pulgadas)
    section = doc.sections[0]
    page_height = section.page_height / EMU_INCH
    page_width = section.page_width / EMU_INCH
    top_margin = section.top_margin / EMU_INCH
    bottom_margin = section.bottom_margin / EMU_INCH
    left_margin = section.left_margin / EMU_INCH
    right_margin = section.right_margin / EMU_INCH

    # Ancho útil de página
    useful_width = page_width - left_margin - right_margin

    # Calcular espacio vertical total usado
    total_vertical_space_used = 0.0
    keys = ['title', 'subtitle', 'description', 'author', 'month', 'year']

    for key in keys:
        text = metadata_dict.get(key, "")
        length = len(text)
        font_size = font_sizes.get(key, DEFAULT_LINE_HEIGHT)
        text_width = length * font_size * CHAR_WIDTH_FACTOR
        lines_used = int(np.ceil(text_width / useful_width)) if useful_width > 0 else 1
        # Espacio vertical: líneas * altura de línea + interlineado extra
        space_used = lines_used * font_size + max(0, lines_used - 1) * font_size
        total_vertical_space_used += space_used

    # Espacio restante y líneas vacías
    useful_height = page_height - top_margin - bottom_margin
    remaining_space = (useful_height - total_vertical_space_used) / 2
    empty_lines = np.ceil(remaining_space / (2 * DEFAULT_LINE_HEIGHT))  # Manteniendo la lógica original

    logger.info(f"DOCX: Vertical centering calculated.")

    return int(empty_lines)

def build_docx_from_dict(doc_dict, buffer, ctx, URL):
    logger.info("DOCX: Starting document generation ...")

    metadata_data = doc_dict.get("metadata", {})
    sections_data = doc_dict.get("sections", [])
    font = doc_dict.get("font", "Times New Roman")
    columns_body = doc_dict.get("columns_body", 1)
    columns_body = int(columns_body)
    if columns_body > 2:
        columns_body = 2  # Limitar a máximo 2 columnas
    elif columns_body < 1:
        columns_body = 1  # Mínimo 1 columna

    doc = Document()

    if str(metadata_data.get("page_break", False)).lower() == "true":
        empty_lines = vertical_center(metadata_data, doc)

    # Parrafos para centrar verticalmente no están soportados en python-docx/DOCX
    if str(metadata_data.get("page_break", False)).lower() == "true":
        for _ in range(empty_lines):
            doc.add_paragraph("")
    
    # Aplicar metadata como portada centrada horizontalmente
    meta = metadata_data
    if "title" in meta:
        title = doc.add_paragraph(meta["title"])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].bold = True
            title.runs[0].font.size = Inches(0.5)  # Tamaño grande para título
            title.runs[0].font.name = font  # Usar font global
    if "subtitle" in meta:
        subtitle = doc.add_paragraph(meta["subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle.runs:
            subtitle.runs[0].italic = True
            subtitle.runs[0].font.size = Inches(0.1667)  # Tamaño mediano para subtítulo
            subtitle.runs[0].font.name = font
    if "description" in meta:
        desc = doc.add_paragraph(meta["description"])
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if desc.runs:
            desc.runs[0].font.size = Inches(0.1667)  
            desc.runs[0].font.name = font
    if "author" in meta:
        author = doc.add_paragraph(f"Autor: {meta['author']}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if author.runs:
            author.runs[0].font.size = Inches(0.1667)
            author.runs[0].font.name = font
    if "month" in meta and "year" in meta:
        date = doc.add_paragraph(f"{meta['month']} {meta['year']}")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if date.runs:
            date.runs[0].font.size = Inches(0.1667)
            date.runs[0].font.name = font
    # Nota: La alineación vertical de página no está soportada en python-docx/DOCX
    # Gestionar salto de página y columnas
    page_break_requested = str(metadata_data.get("page_break", False)).lower() == "true"

    if columns_body > 1:
        # Si hay columnas, usamos el section break para controlar el salto de página
        section_type = WD_SECTION_START.NEW_PAGE if page_break_requested else WD_SECTION_START.CONTINUOUS
        new_section = doc.add_section(start_type=section_type)
        sectPr = new_section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), str(columns_body))
        sectPr.append(cols)
    elif page_break_requested:
        # Si no hay columnas pero sí salto, usamos salto manual
        doc.add_page_break()
    
    # Contadores para numeración automática de captions
    figure_counter = 1
    table_counter = 1
    equation_counter = 1
    
    # Variable para acumular párrafos continuos
    current_paragraph = None
    
    # Procesar elementos
    for item in sections_data:
        # Detect type using the discriminator first if available
        item_type = item.get("type", None)
        
        if item_type == "ParagraphHeader" or ("text" in item and "level" in item):  # ParagraphHeader
            current_paragraph = None  # Reset paragraph
            heading = doc.add_heading(item["text"], level=item.get("level", 2))
            # if item.get("alignment") == "center":
            #     heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if heading.runs:
                heading.runs[0].font.name = font
        
        elif item_type == "ParagraphBody" or (item_type is None and "text" in item and "bold" not in item):  # ParagraphBody
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
                # Set alignment for the paragraph based on the first ParagraphBody
                # if item.get("alignment") == "center":
                #     current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # elif item.get("alignment") == "justify":
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                # Default is left
            # Add run to the current paragraph
            run = current_paragraph.add_run(item.get("text", ""))
            # ParagraphBody shouldn't have bold/italic flags in the new schema, but handle gently if present
            run.font.size = Inches(12 / 72)
            run.font.name = font

        elif item_type == "WordsWithBoldOrItalic" or ("text" in item and ("bold" in item or "italic" in item)): # WordsWithBoldOrItalic
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            
            run = current_paragraph.add_run(item.get("text", ""))
            run.bold = item.get("bold", False)
            run.italic = item.get("italic", False)
            run.font.size = Inches(12 / 72)
            run.font.name = font
        
        elif item_type == "ParagraphListItem" or "items" in item:  # ParagraphListItem
            current_paragraph = None  # Reset paragraph
            for it in item["items"]:
                p = doc.add_paragraph(it, style='List Bullet' if item.get("list_style") == "List Bullet" else 'List Number')
                if p.runs:
                    p.runs[0].font.name = font
        
        elif item_type == "Table" or "headers" in item:  # Table
            current_paragraph = None  # Reset paragraph
            if not item.get("headers") or not item.get("rows"):
                raise ValueError("Table must have headers and rows.")
            caption_text = item.get("caption", f"Table {table_counter}: ")
            if not item.get("caption"):
                table_counter += 1
            p = doc.add_paragraph(caption_text, style='Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=1, cols=len(item["headers"]))
            table.style = 'Light List Accent 1'
            hdr_cells = table.rows[0].cells
            for i, hdr in enumerate(item["headers"]):
                hdr_cells[i].text = hdr
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.name = font
            for row_data in item["rows"]:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data
                    for run in row_cells[i].paragraphs[0].runs:
                        run.font.name = font
        
        elif item_type == "Image" or "id" in item:  # Image
            current_paragraph = None  # Reset paragraph
            try:
                bearer_token = _get_bearer_token(ctx)
                image_file = download_file(URL, bearer_token, item["id"])
                if isinstance(image_file, dict) and "error" in image_file:
                    raise ValueError(f"Error downloading image with ID {item['id']}: {image_file['error']['message']}")
                if not image_file or (hasattr(image_file, 'getbuffer') and len(image_file.getbuffer()) == 0):
                    raise ValueError(f"Downloaded image with ID {item['id']} is empty or invalid.")
                img_width, img_height = img_dimensions(image_file, body_columns=columns_body)
                doc.add_picture(image_file, width=Inches(img_width), height=Inches(img_height))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f"Failed to load image {item['id']}: {e}. Adding placeholder.")
                doc.add_paragraph(f"[Image Placeholder: {item.get('caption', 'No caption')}]", style='Caption')
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_text = item.get("caption", f"Figure {figure_counter}: ")
            if not item.get("caption"):
                figure_counter += 1
            p = doc.add_paragraph(caption_text, style='Caption')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif item_type == "Equation" or "latex" in item:  # Equation
            current_paragraph = None  # Reset paragraph
            p = doc.add_paragraph()
            math2docx.add_math(p, item["latex"])
            if item.get("caption"):
                caption_text = item.get("caption", f"Equation {equation_counter}: ")
                if not item.get("caption"):
                    equation_counter += 1
                p_cap = doc.add_paragraph(caption_text, style='Caption')
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        else:
            current_paragraph = None  # Reset for any other item

    logger.info("DOCX: Document generation completed!")
    doc.save(buffer)
    buffer.seek(0)
    return buffer